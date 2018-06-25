#!/usr/bin/env python3.6
# -*- coding: utf-8 -*-
"""
:DESCRIPTION:
    - The purpose of this script is to break up the Kaggle training set and
        validation set data into individual .docx and .txt documents to
        prove the validity of our proposed ingest capabilities.
    - This script will create duplicates copies (1 .txt and 1 .docx) for every
        row in the .xlsx files.

:REQUIRES:
    - Python-Docx package (https://python-docx.readthedocs.io/en/latest/)
        * install via the command:  'pip install python-docx'
        * call library via the command: 'import docx'

:TODO:
    NONE
"""

#==============================================================================
# Package Import
#==============================================================================
import os
import pandas as pd
import docx
from pathlib import Path

#==============================================================================
# Function Definitions / Reference Variable Declaration
#==============================================================================
def convert2docx(data, col_names, setnum):
    '''
    Description:
        - This function will take the passed in string from a row in a .xlsx
            or .csv file and create a new .docx file containing only this string
    Input:
        - data (STRING):  a string containing the full contents of an essay
            submission along with any additional scoring information for that
            essay

    Output:
        - a .docx file will be written to the /Data/txt subfolder within the
            main project working folder
    '''
    # create the .docx file
    doc = docx.Document()

    # output the contents of each column to the .docx file
    for i in range(0,len(data)):
        doc.add_heading(col_names[i], 2)
        doc.add_paragraph(str(data[i]))

    # check to make sure a folder exists in the target directory for the set
    dir_path = Path('Data',setnum)
    if os.path.isdir(dir_path) == False:
        os.makedirs(dir_path)

    # check to make sure a `docx` folder exists in the target directory
    dir_path = Path(dir_path, 'docx')
    if os.path.isdir(Path(dir_path)) == False:
        os.makedirs(Path(dir_path))

    # save the created .docx file
    doc.save(str(Path(dir_path, str(data['essay_id']) + '.docx')))

def convert2txt(data, col_names, setnum):
    '''
    Description:
        - This function will take the passed in string from a row in a .xlsx
            or .csv file and create a new .txt file containing only this string

    Input:
        - data (STRING):  a string containing the full contents of an essay
            submission along with any additional scoring information for that
            essay

    Output:
        - a .txt file will be written to the /Data/txt subfolder within the
            main project working folder
    '''
    # check to make sure a folder exists in the target directory for the set
    dir_path = Path('Data',setnum)
    if os.path.isdir(dir_path) == False:
        os.makedirs(dir_path)

    # check to make sure a `txt` folder exists in the target directory
    if os.path.isdir(Path(dir_path, 'txt')) == False:
        os.makedirs(Path(dir_path, 'txt'))

    filename = str(Path(dir_path, 'txt', str(data['essay_id']) + '.txt'))
    with open(filename, "w", encoding = 'utf-8') as my_output_file:
        for i in range(0, len(data)):
            my_output_file.write(col_names[i] + ": " + str(data[i]) + '\n')
    my_output_file.close()

def unpack_training_files():
    '''
    Description:
        - This is function is custom built for the current prototype
        version of SentNet. It will ingest the .csv file containing all the
        essays we will be testing SentNet against and unpack them into
        individual documents (.docx).

    Input:
        - NONE
        - the .csv file is hard coded into this function:
            `training_set_rel3.csv` located in the `Data` folder of the SentNet
            directory

    Output:
        - .docx files are created for every row in the original .csv file
            * documents are contained within subfolders based on their essay
                number (e.g., all essays from Set 4 will be in the `Set 4`
                folder)
        - subsets of the original .csv file are also created that contain
            only the essays within a specific set should future developers
            or analysts wish to examine the original data for a particular set
    '''
    df = pd.read_csv(Path('Data','training_set_rel3.csv'), encoding = 'windows-1252')

    # For every essay set in the data, extract the data for that set
    #   and create an individual file for every row in that set
    essay_set_list = df['essay_set'].value_counts().sort_index().index.tolist()
    for df_num in essay_set_list:
        # create the subset of the original data based on the `essay_id`
        df_sub = df[df['essay_set'] == df_num]
        df_sub.reset_index(inplace=True)
        del df_sub['index']
        df_sub = df_sub.dropna(thresh=0.8*len(df_sub), axis=1)

        # create file names and folder names for the new data
        file = 'Set' + str(df_num)
        filename = (str(Path('Data','Set' + str(df_num), str(
                'training_set_rel3_set' + str(df_num) + '.csv'))))

        # create the individual .docx and .txt files for every row in the subset
        columns = list(df_sub.columns)
        for index, row in df_sub.iterrows():
            convert2docx(row, columns, file)
            # disabled .txt output at this time as not used in prototype
            #convert2txt(row, columns, file)
            if index%100 == 0:
                print('Done unpacking: ' + str(index) + ' files from ' + file)

        # create a .csv of the subset dataframe w/ only the essays in that id group
        df_sub.to_csv(filename, encoding = 'utf-8-sig', index='False')

#==============================================================================
# Working Code
#==============================================================================

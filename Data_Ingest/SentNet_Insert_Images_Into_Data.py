#!/usr/bin/env python3.6
# -*- coding: utf-8 -*-
"""
Created on Sun Jun 10 18:23:08 2018

@author: Eric Reidelbach

:DESCRIPTION:
    - This script will inject report-related images (i.e. bar charts, line
        charts) into our training and test data for the purposes of simulating
        visualizations in intelligence reports. The function will randomly pull
        from a pool of images and insert 0, 1, 2 images (# is chosen at
        random) and insert the images into a file.
    - There are three different image types we'll insert:
        * bar chart, pie chart and/or line chart
    - Although SentNet has the capability of ingesting .docx AND .txt files,
        this script will only attempt to insert images or image related data
        into .docx files.

:REQUIRES:
    - Python-Docx package (https://python-docx.readthedocs.io/en/latest/)
        * install via the command:  'pip install python-docx'
        * call library via the command: 'import docx'
    - docx2txt package (https://github.com/ankushshah89/python-docx2txt)
        * install via the command: 'pip install docx2txt'
        * import library via the command:  'import docx2txt'
    
:TODO:
"""
 
#==============================================================================
# Package Import
#==============================================================================
import docx2txt
import docx
import os   
import random
from Data_Ingest.SentNet_Docx_Txt_Ingest import ingest_files

#==============================================================================
# Function Definitions / Reference Variable Declaration
#==============================================================================
def insert_image_into_docx(doc_dir_path, img_dir_path):
    '''
   Description:
       - This function will ingest every .docx file contained within the 
           specified folder (i.e. `path` )and inject 0-3 images into the file.
           The number is randomly chosen.
    - There are three image types in the directory: Pie Charts, Bar Charts,
        and Line Charts with 10 images in each type (30 total)
    - Images will be distributed by score such that:
        * bottom 30% of `domain1_score` will pull from images 1, 2, 3, or 4
        * middle 40% of `domain1_score will pull from images 3, 4, 5, 6, 7 or 8
        * top 30% of `domain1_score will pull from images 7, 8, 9, or 10
    Input:
        - doc_dir_path (STRING):  the directory that contains the .docx data
        - img_dir_path (STRING): the directory containing all images that
            will be chosen from to be inserted into a .docx file
    Output:
        - Nothing is returned by this function
        - This function will overwrite every .docx file with a new file
            containing 0-3 images
    '''
    
    # setup variables for use in the function
    imageTypeDict = {'0':[],
                 '1':['PieChart'],
                 '2':['LineChart'],
                 '3':['BarChart'],
                 '4':['PieChart','LineChart'],
                 '5':['PieChart','BarChart'],
                 '6':['BarChart','LineChart'],
                 '7':['PieChart','LineChart','BarChart'],
                 }
    
    # read in the filename of each image in the image directory into a list
    image_list = []
    for fname in os.listdir(img_dir_path):
        image_list.append(os.path.join(img_dir_path, fname))
        
    # read in every .docx file in the specified path: `doc_dir_path
    doc_list = []
    file_list = []
    for filename in os.listdir(doc_dir_path):
        if filename.endswith('.docx'):
            file_list.append(filename)
            doc_list.append(os.path.join(doc_dir_path, filename))
        
    # read in the contents of each file into a dataframe so that we can compute
    #   the distribution of `domain_scores`
    docDF = ingest_files(doc_dir_path)
    # extract the scores, convert them to integers and sort them
    domain_scores = docDF['domain1_score'].value_counts().index.tolist()
    domain_scores = sorted([int(x.split('.')[0]) for x in domain_scores])
        
    # iterate over every file in the list
    for file_name, file_path, in zip(file_list, doc_list):
        # check to see what the file's value is for `domain1_score`
        doc = docx2txt.process(file_path)
        
        # split the document into a list and then extract the `domain1_score`
        doc_elements = doc.split('\n\n')
        domScore = int(doc_elements[doc_elements.index('domain1_score')+1].split('.')[0])
        
        # randomly select what image types should be put in the file:
        #   0, 1, 2, or 3 of the available options:  pie, line or bar
        imageTypeList = imageTypeDict[str(random.randint(0,7))]
        
        # randomly select what images should be pulled from the group based on the
        #   file's `domain1_score`
        insertedImagesList = []
        percentile = (domain_scores.index(domScore)+1)/len(domain_scores)
        for image in imageTypeList:
            if percentile <= 0.3:   # bottom 30%
                insertedImagesList.append(image + str(random.randint(1,4)) + '.png')
            elif percentile > 0.3 and percentile < 0.7: # middle 40%
                insertedImagesList.append(image + str(random.randint(3,8)) + '.png')
            else:   # upper 30%
                insertedImagesList.append(image + str(random.randint(7,10)) + '.png')
                      
        # open the specified .docx document for editing
        document = docx.Document(file_path)
        
        # inject the image(s) into the file under the header(s): Image1, Image2, etc.
        for image in insertedImagesList:
            try:
                image_path = [x for x in image_list if image in x][0]
            except:
                print(image)
            heading = 'Image ' + str(insertedImagesList.index(image)+1)
            document.add_heading(heading, 2)
            document.add_picture(image_path)
    
        # check to make sure an image subfolder exists in the docx folder
        # if it doesn't, make it
        if os.path.isdir(os.path.join(file_path.split(
                '\docx')[0], 'docx', 'Images')) == False:
            os.makedirs(os.path.join(file_path.split(
                    '\docx')[0], 'docx', 'Images'))
        #os.chdir(os.path.join(file_path.split('\docx')[0], 'docx', 'Images'))
        
        # save the file
        document.save(os.path.join(file_path.split(
                '\docx')[0],'docx', 'Images',file_name))
        
        # keep a counter going for progress
        if file_list.index(file_name)%100 == 0:
            print('Done with ' + str(file_list.index(file_name)) + ' files.')
        
#==============================================================================
# Working Code
#==============================================================================

# Set the project working directory
os.chdir(r'E:\Projects\SentNet\Data_Ingest')

# Inject images into every .docx file for Set7
insert_image_into_docx('E:\Projects\SentNet\Data\Set7\docx',
                       'E:\Projects\SentNet\Data\Images')

# Inject images into every .docx file for Set8
insert_image_into_docx('E:\Projects\SentNet\Data\Set7\docx',
                       'E:\Projects\SentNet\Data\Images')